from docxtpl import DocxTemplate
from docx import Document

doc = DocxTemplate("testDoc.docx")

customerNames = []
customerNames.append({"FirstName":'james' ,
                "LastName":'don' ,
                "Description":'yes'})

customerNames.append({"FirstName":'chris' ,
                "LastName":'witt' ,
                "Description":'dklfdj'})

customerNames.append({"FirstName":'sleepy' ,
                "LastName":'joe' ,
                "Description":'dfdf'
                })
customerNames.append({"FirstName":'hohn' ,
                "LastName": '2.0',
                "Description":'dfdfd'})

newName = [
  {
    "FirstName":"NewFirst1",
    "LastName":'NewLast1',
    "Description":"NewDescription1"
  },
  {
    "FirstName":"NewFirst2",
    "LastName":'NewLast2',
    "Description":"NewDescription2"
  },
  {
    "FirstName":"NewFirst3",
    "LastName":'NewLast3',
    "Description":"NewDescription3"
  },
  {
    "FirstName":"NewFirst4",
    "LastName":'NewLast4',
    "Description":"4"
  },
]
Cars =[
{'Make':'Ford',
'Model':'Falcon',
'Color':'White'},

{'Make':'Holden',
'Model':'Commador',
'Color':'Black'},

{'Make':'Nissan',
'Model':'Skyline',
'Color':'Blue'},

{'Make':'Toyota',
'Model':'Carola',
'Color':'Red'}
]
topItemsRow = ['coke','tv','battery']

context = {'Date' : "Friday Bitches",
            "customerNames": customerNames,
            "newNames":newName,
            'Cars':Cars
}
# print(customerNames[0])
print(newName[0].items)
doc.render(context)
doc.save("generated_doc.docx")

document = Document('./generated_doc.docx')
print(document.tables[2])
table3 = document.tables[2]._cells
for item in (table3):
  print(item.text)

allTables = document.tables
for activeTable in allTables:
  value = (activeTable.cell(1,2).paragraphs[0].text)
  print(activeTable.cell(1,2).paragraphs[0].text)
  if value == 'don':
    print('removing table')
    activeTable._element.getparent().remove(activeTable._element)

document.save('generated2.docx')